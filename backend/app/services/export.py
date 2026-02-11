import io
import json
import uuid

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.errors import NotFoundError
from app.models.bank_instance import BankInstance
from app.models.perspective import Perspective
from app.models.vdba import Vdba
from app.services.minio import ensure_bucket, get_minio_client


async def generate_export(db: AsyncSession, vdba_id: uuid.UUID) -> str:
    """Generate an export file for a VDBA and upload to MinIO.

    Supports JSON, PDF, and DOCX formats with fallback to JSON when
    optional libraries (reportlab, python-docx) are not installed.

    Returns the MinIO key of the uploaded file.
    """
    result = await db.execute(select(Vdba).where(Vdba.id == vdba_id))
    vdba = result.scalar_one_or_none()
    if not vdba:
        raise NotFoundError("VDBA not found")

    # Gather all bank instances for the journey
    bank_result = await db.execute(
        select(BankInstance)
        .join(Perspective, BankInstance.perspective_id == Perspective.id)
        .where(Perspective.journey_id == vdba.journey_id)
        .order_by(BankInstance.created_at)
    )
    bank_instances = list(bank_result.scalars().all())

    # Gather perspectives for context
    persp_result = await db.execute(
        select(Perspective)
        .where(Perspective.journey_id == vdba.journey_id)
        .order_by(Perspective.dimension, Perspective.phase)
    )
    perspectives = list(persp_result.scalars().all())

    # Build export data structure
    export_data = _build_export_data(vdba, perspectives, bank_instances)

    # Generate file based on format
    if vdba.export_format == "json":
        file_bytes, content_type, extension = _generate_json(export_data)
    elif vdba.export_format == "pdf":
        file_bytes, content_type, extension = _generate_pdf(export_data)
    elif vdba.export_format == "docx":
        file_bytes, content_type, extension = _generate_docx(export_data)
    else:
        file_bytes, content_type, extension = _generate_json(export_data)

    # Upload to MinIO
    client = get_minio_client()
    await ensure_bucket(client)

    minio_key = f"exports/vdba/{vdba.id}.{extension}"
    data_stream = io.BytesIO(file_bytes)
    from app.core.config import settings

    await client.put_object(
        settings.minio_bucket,
        minio_key,
        data_stream,
        length=len(file_bytes),
        content_type=content_type,
    )

    # Update VDBA export_url
    vdba.export_url = minio_key
    await db.flush()

    return minio_key


def _build_export_data(
    vdba: Vdba,
    perspectives: list[Perspective],
    bank_instances: list[BankInstance],
) -> dict:
    """Build a structured dictionary of all VDBA data for export."""
    return {
        "vdba": {
            "id": str(vdba.id),
            "title": vdba.title,
            "description": vdba.description,
            "published_at": vdba.published_at.isoformat() if vdba.published_at else None,
            "export_format": vdba.export_format,
            "version": vdba.version,
        },
        "perspectives": [
            {
                "id": str(p.id),
                "dimension": p.dimension,
                "phase": p.phase,
                "status": p.status,
                "started_at": p.started_at.isoformat() if p.started_at else None,
                "completed_at": p.completed_at.isoformat() if p.completed_at else None,
            }
            for p in perspectives
        ],
        "bank_instances": [
            {
                "id": str(b.id),
                "type": b.type,
                "synopsis": b.synopsis,
                "agent_assessments": b.agent_assessments,
                "decision_audit": b.decision_audit,
                "documents_count": b.documents_count,
                "vibes_count": b.vibes_count,
            }
            for b in bank_instances
        ],
    }


def _generate_json(data: dict) -> tuple[bytes, str, str]:
    """Generate JSON export."""
    content = json.dumps(data, indent=2, default=str)
    return content.encode("utf-8"), "application/json", "json"


def _generate_pdf(data: dict) -> tuple[bytes, str, str]:
    """Generate PDF export. Falls back to JSON if reportlab is not installed."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        # Title
        vdba_info = data["vdba"]
        story.append(Paragraph(f"VDBA: {vdba_info['title']}", styles["Title"]))
        story.append(Spacer(1, 12))

        if vdba_info.get("description"):
            story.append(Paragraph(vdba_info["description"], styles["Normal"]))
            story.append(Spacer(1, 12))

        story.append(Paragraph(f"Version: {vdba_info['version']}", styles["Normal"]))
        story.append(Paragraph(f"Published: {vdba_info['published_at']}", styles["Normal"]))
        story.append(Spacer(1, 20))

        # Perspectives
        story.append(Paragraph("Perspectives", styles["Heading2"]))
        for p in data["perspectives"]:
            story.append(
                Paragraph(
                    f"{p['dimension'].title()} - {p['phase'].title()}: {p['status']}",
                    styles["Normal"],
                )
            )
        story.append(Spacer(1, 20))

        # Bank Instances
        story.append(Paragraph("Bank Instances", styles["Heading2"]))
        for b in data["bank_instances"]:
            story.append(Paragraph(f"Type: {b['type']}", styles["Heading3"]))
            if b.get("synopsis"):
                synopsis_text = b["synopsis"][:500]
                story.append(Paragraph(f"Synopsis: {synopsis_text}", styles["Normal"]))
            story.append(Spacer(1, 8))

        doc.build(story)
        return buffer.getvalue(), "application/pdf", "pdf"
    except ImportError:
        # reportlab not installed, fall back to JSON
        return _generate_json(data)


def _generate_docx(data: dict) -> tuple[bytes, str, str]:
    """Generate DOCX export. Falls back to JSON if python-docx is not installed."""
    try:
        from docx import Document

        doc = Document()

        vdba_info = data["vdba"]
        doc.add_heading(f"VDBA: {vdba_info['title']}", level=0)

        if vdba_info.get("description"):
            doc.add_paragraph(vdba_info["description"])

        doc.add_paragraph(f"Version: {vdba_info['version']}")
        doc.add_paragraph(f"Published: {vdba_info['published_at']}")

        # Perspectives
        doc.add_heading("Perspectives", level=1)
        for p in data["perspectives"]:
            doc.add_paragraph(
                f"{p['dimension'].title()} - {p['phase'].title()}: {p['status']}"
            )

        # Bank Instances
        doc.add_heading("Bank Instances", level=1)
        for b in data["bank_instances"]:
            doc.add_heading(f"Type: {b['type']}", level=2)
            if b.get("synopsis"):
                doc.add_paragraph(f"Synopsis: {b['synopsis'][:500]}")

        buffer = io.BytesIO()
        doc.save(buffer)
        return (
            buffer.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx",
        )
    except ImportError:
        # python-docx not installed, fall back to JSON
        return _generate_json(data)
